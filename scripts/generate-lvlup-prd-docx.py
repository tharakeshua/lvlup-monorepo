#!/usr/bin/env python3
"""Generate LVLUP-PRODUCT-REQUIREMENTS.docx from journey and requirements sources."""

from __future__ import annotations

import csv
import re
import sys
from datetime import date
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "requirements" / "LVLUP-PRODUCT-REQUIREMENTS.docx"
TMP_MD = ROOT / "docs" / "requirements" / "_LVLUP-PRODUCT-REQUIREMENTS-combined.md"
TMP_DOCX = ROOT / "docs" / "requirements" / "_LVLUP-PRODUCT-REQUIREMENTS-body.docx"

JOURNEYS = ROOT / "tmp" / "pr-factory-backup" / "docs" / "journeys"
EXAM_JOURNEY = ROOT / "docs" / "requirements" / "EXAM-QR-BATCH-JOURNEY.md"
TRACKER_CSV = ROOT / "docs" / "requirements" / "EXAM-QR-BATCH-JOURNEY-TRACKER.csv"
VERSION = "1.0"
DOC_DATE = "2026-07-14"


def read_text(path: Path) -> str:
    if not path.exists():
        return f"_Source not found: `{path.relative_to(ROOT)}`_\n"
    return path.read_text(encoding="utf-8")


def strip_doc_header(md: str) -> str:
    """Remove duplicate top-level title and metadata block from child docs."""
    lines = md.splitlines()
    out: list[str] = []
    skip_meta = False
    for i, line in enumerate(lines):
        if i == 0 and line.startswith("# "):
            continue
        if line.strip() == "---" and i < 20:
            skip_meta = not skip_meta
            continue
        if skip_meta:
            continue
        out.append(line)
    return "\n".join(out).strip()


def bump_headings(md: str, levels: int = 1) -> str:
    """Increase markdown heading depth by `levels`."""
    if levels <= 0:
        return md
    result: list[str] = []
    for line in md.splitlines():
        m = re.match(r"^(#{1,6})\s", line)
        if m:
            hashes = m.group(1)
            new_depth = min(len(hashes) + levels, 6)
            result.append("#" * new_depth + line[len(hashes) :])
        else:
            result.append(line)
    return "\n".join(result)


def load_tracker_table() -> str:
    if not TRACKER_CSV.exists():
        return "_Tracker CSV not found._\n"
    rows: list[dict[str, str]] = []
    with TRACKER_CSV.open(encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)

    lines = [
        "| ID | Phase | Step | Requirement | Status | Priority | Notes |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for r in rows:
        lines.append(
            "| {id} | {phase} | {step} | {req} | {status} | {priority} | {notes} |".format(
                id=r.get("id", ""),
                phase=r.get("phase", ""),
                step=r.get("journey_step", ""),
                req=r.get("requirement", "").replace("|", "\\|"),
                status=r.get("status", ""),
                priority=r.get("priority", ""),
                notes=r.get("notes", "").replace("|", "\\|"),
            )
        )
    return "\n".join(lines) + "\n"


def gap_analysis_section() -> str:
    return """## Gap Analysis Summary

| Stage / Capability | Status | Notes |
| --- | --- | --- |
| Teacher login (GRN001 flow) | **EXISTS** | School code + email/password |
| Exam create + QP upload + AI extract | **EXISTS** | teacher-web wizard |
| Manual offline exam (paper) | **Process** | Outside software |
| Teacher-web answer sheet upload | **EXISTS** | Submissions page |
| Real-time AI grading pipeline | **EXISTS** | `advancePipeline` |
| Teacher review + manual override | **EXISTS** | GradingReviewPage |
| Results release + parent PDF | **PARTIAL** | PDF exists; release gate gaps |
| Student results (print, not PDF) | **PARTIAL** | `window.print()` only |
| Parent in-app notification | **PARTIAL** | Legacy trigger; v1 outbox gap |
| Post-exam auto-space | **GAP** | Not implemented |
| QR Agent PWA (`apps/scanner-web`) | **GAP** | Backend seams exist |
| Handwriting anti-fraud match | **GAP** | P0 security requirement |
| EOD batch grading scheduler | **GAP** | 15-min watchdog only |
| Scanner role backend (`uploadAnswerSheets`) | **EXISTS** | No frontend app |
| Parent PDF generateReport | **EXISTS** | ExamResultsPage |
| LD-01 remediation content tests | **PARTIAL** | Convention not type |
| Legacy autograde triggers | **PARTIAL** | Dual deployment |

### Tracker Detail (EXAM-QR-BATCH-JOURNEY-TRACKER.csv)

""" + load_tracker_table()


def implementation_phases() -> str:
    return """## Implementation Phases

### P0 — Must ship for MVP journey

- [ ] **P0-01** Scaffold `apps/scanner-web` with login, exam select, QR scan, capture, upload
- [ ] **P0-02** `generateExamQrBatch` + `validateQrScan` callables + `examQrTokens` collection
- [ ] **P0-03** Handwriting verification service with cross-student block (HW-04, HW-06)
- [ ] **P0-04** Agent session end + QR printed/scanned audit
- [ ] **P0-05** `eodGradingBatch` scheduler (modified-date trigger)
- [ ] **P0-06** Fix v1 `releaseResults` outbox recipient resolution
- [ ] **P0-07** Student PDF download (parity with parent)
- [ ] **P0-08** `generateReport` release gate
- [ ] **P0-09** LD-01 regression tests on remediation content
- [ ] **P0-10** End-to-end Playwright: exam create → scanner upload → batch → release → student results

### P1 — Quality + automation

- [ ] **P1-01** Auto-create remediation space on `results_released`
- [ ] **P1-02** `extractConceptsFromExam` from QP
- [ ] **P1-03** Wrong-answer practice mode routed from `ExamResultPage`
- [ ] **P1-04** FCM push for parent results notification
- [ ] **P1-05** `saveExamQuestion` callable for rubric persistence
- [ ] **P1-06** Scanner offline queue (IndexedDB)
- [ ] **P1-07** Deprecate legacy `functions/autograde` triggers (migrate to sdk-v1)
- [ ] **P1-08** Admin UI: scanner account provisioning + QR batch print

### P2 — Polish

- [ ] **P2-01** Per-concept retest story point type
- [ ] **P2-02** Email PDF delivery on release
- [ ] **P2-03** Native scanner app evaluation
- [ ] **P2-04** Multi-language QR payload
- [ ] **P2-05** Brush-up viewer from QP PDF annotations
"""


def appendix_section() -> str:
    return """## Appendix

### Demo Credentials (pointer only — no secrets in this document)

All demo emails, passwords, and school codes are documented in **`TEST_CREDENTIALS.md`** at the repository root.

| School code | Purpose |
| --- | --- |
| `GRN001` | Greenwood demo tenant (primary QA) |
| `SUB001` | Subhang demo tenant |
| `SPR001` | Emulator Springfield tenant |

Example accounts (password pattern documented in TEST_CREDENTIALS.md):

| Role | Example email | School code |
| --- | --- | --- |
| Teacher | `priya.sharma@greenwood.edu` | GRN001 |
| Admin | `admin@greenwood.edu` | GRN001 |
| Student | `aarav.patel@greenwood.edu` or roll `2025001` | GRN001 |
| Parent | `suresh.patel@gmail.com` | GRN001 |
| Super-admin | `superadmin@levelup.app` | _(no school code)_ |
| Consumer (B2C) | `consumer@gmail.test` | _(no school code)_ |

Seed scripts: `pnpm seed:production` · emulator `pnpm seed:emulator`.

### Local Development Ports

| App | Path | Dev port | Hosting target |
| --- | --- | --- | --- |
| Super-admin | `apps/super-admin` | **4567** | `super-admin` |
| Admin | `apps/admin-web` | **4568** | `admin-web` |
| Teacher | `apps/teacher-web` | **4569** | `teacher-web` |
| Student | `apps/student-web` | **4570** | `student-web` |
| Parent | `apps/parent-web` | **4571** | `parent-web` |
| Scanner PWA (planned) | `apps/scanner-web` | **4574** | `scanner-web` |
| Marketing website | `apps/website` | **4321** | `website` |

### Firebase Emulator Ports

| Service | Port |
| --- | --- |
| Emulator UI | 4000 |
| Auth | 9099 |
| Firestore | 8080 |
| Functions | 5001 |
| RTDB | 9000 |

### Tech Stack

| Layer | Technology |
| --- | --- |
| Frontend | React 18, Vite, Tailwind CSS 3.4, shadcn/ui, Radix UI |
| Backend | Firebase Cloud Functions (sdk-v1 + legacy codebases) |
| Database | Firestore (`v2_*` prefix via `LVLUP_COLLECTION_PREFIX`) |
| Real-time | Firebase RTDB (leaderboards, grading progress) |
| Storage | Firebase Storage (question papers, answer sheets, reports) |
| Auth | Firebase Auth with custom JWT claims |
| AI | Google Gemini (`gemini-2.5-pro`, `gemini-2.5-flash`) via `@levelup/ai` |
| Async | Cloud Tasks (`advancePipeline`) |
| Hosting | Firebase Hosting (6 targets) |
| Monorepo | pnpm workspaces, Turborepo |
| E2E | Playwright |
| Unit tests | Vitest |

### Environment / Secret Names (pointers only)

| Name | Purpose |
| --- | --- |
| `VITE_FIREBASE_*` | Client Firebase config (no AI keys in client) |
| `LVLUP_COLLECTION_PREFIX` | Production collection prefix (`v2_`) |
| `LEVELUP_AI_KEY` / `GEMINI_API_KEY` | Dev/emulator AI override |
| `LEVELUP_AI_MODEL_PRO` / `LEVELUP_AI_MODEL_FLASH` | Model IDs |
| `tenant-{tenantId}-gemini` | Secret Manager per-tenant Gemini key |
| `GOOGLE_CLOUD_PROJECT` / `GCLOUD_PROJECT` | GCP project for Secret Manager |
| `GOOGLE_APPLICATION_CREDENTIALS` | Service account path for local functions (never commit) |

### Related Documents

- `docs/requirements/EXAM-QR-BATCH-JOURNEY.md` — P0 journey requirements
- `docs/requirements/EXAM-QR-BATCH-JOURNEY-TRACKER.csv` — Implementation tracker
- `tmp/pr-factory-backup/docs/journeys/` — Role journey guides (01–06)
- `docs/handover/QA-ASSIGN-TAKE-FLOW.md` — Assign/take handover
- `TEST_CREDENTIALS.md` — Demo credentials (repo root)
- `docs/rebuild-spec/status/SDK-REVIEW-A-LEARNING-DOMAIN.md` — LD-01 answer key leak spec
"""


def build_combined_markdown() -> str:
    guide = read_text(JOURNEYS / "LVLUP-JOURNEY-GUIDE.md")
    student = bump_headings(strip_doc_header(read_text(JOURNEYS / "01-student-journey.md")), 1)
    teacher = bump_headings(strip_doc_header(read_text(JOURNEYS / "02-teacher-journey.md")), 1)
    parent = bump_headings(strip_doc_header(read_text(JOURNEYS / "03-parent-journey.md")), 1)
    admin = bump_headings(strip_doc_header(read_text(JOURNEYS / "04-admin-journey.md")), 1)
    cross = bump_headings(strip_doc_header(read_text(JOURNEYS / "05-cross-role-and-ai.md")), 1)
    infra = bump_headings(strip_doc_header(read_text(JOURNEYS / "06-marketing-infra-firebase-vercel-railway.md")), 1)
    exam = read_text(EXAM_JOURNEY)

    # Extract data model section from exam journey for dedicated chapter
    data_model_match = re.search(
        r"(## Data Model / Schema.*?)(?=\n## API / Callables Inventory)",
        exam,
        re.DOTALL,
    )
    data_model = data_model_match.group(1).strip() if data_model_match else ""

    # Exam journey chapters 1-8 as P0 section (full exam doc minus duplicate exec summary at top)
    exam_body = re.sub(
        r"^# Exam → QR Agent.*?^---\n\n## Executive Summary.*?\n---\n\n",
        "",
        exam,
        count=1,
        flags=re.DOTALL | re.MULTILINE,
    )

    parts = [
        "# LvlUp Product Requirements\n",
        f"**Version:** {VERSION}  \n",
        f"**Date:** {DOC_DATE}  \n",
        f"**Status:** Production requirements  \n",
        f"**Workspace:** startup-mvp/lvlup\n",
        "\n\\newpage\n",
        "# Executive Summary\n",
        strip_doc_header(guide).split("## 2. Cross-role journey")[0].replace(
            "# LvlUp Complete Product Journey Guide", ""
        ).strip(),
        "\n\n### Platform Overview\n",
        "LvlUp is a multi-tenant educational platform connecting **offline paper exams** to "
        "**AI-assisted grading** and **personalized digital learning**. The platform serves "
        "schools (B2B) and individual learners (B2C consumer shell on student-web).\n",
        "\n### Personas\n",
        "| Persona | App / Surface | Auth | Primary goals |",
        "| --- | --- | --- | --- |",
        "| **Teacher** | `apps/teacher-web` (:4569) | School code + email + password | Create exams, author spaces, review AI grades, release results |",
        "| **Student** | `apps/student-web` (:4570) | School code + roll/email + password | Learn, practice, take tests, view released exam results |",
        "| **Parent** | `apps/parent-web` (:4571) | School code + email + password | View child results, progress, alerts; download PDF reports |",
        "| **Admin** | `apps/admin-web` (:4568) | School code UX + email/password | Provision users, classes, sessions; school analytics |",
        "| **Manual Agent** | `apps/scanner-web` (planned :4574) | Admin-provisioned scanner credentials | QR print/attach, scan, capture answer sheets |",
        "| **Super-admin** | `apps/super-admin` (:4567) | Email + password + superAdmin claim | Platform tenants, feature flags, LLM usage |",
        "\n\\newpage\n",
        "# Overall Platform Journey\n",
        "## Two Product Tracks\n",
        cross.split("## Two product tracks")[1].split("## End-to-end actor map")[0].strip()
        if "Two product tracks" in cross
        else "",
        "\n\n## End-to-End Actor Map\n",
        cross.split("## End-to-end actor map")[1].split("## Auth / session")[0].strip()
        if "End-to-end actor map" in cross
        else "",
        "\n\n## Cross-Role Sequence (P0 Exam Flow)\n",
        "See P0 Exam-QR-Batch Journey section for the full mermaid sequence diagram.\n",
        "\n\\newpage\n",
        "# Per-Role Journeys\n",
        "## Teacher Journey\n",
        teacher,
        "\n\\newpage\n",
        "## Student Journey\n",
        student,
        "\n\\newpage\n",
        "## Parent Journey\n",
        parent,
        "\n\\newpage\n",
        "## Admin Journey\n",
        admin.split("## B. Super-Admin")[0].strip(),
        "\n\\newpage\n",
        "## Manual Agent (Scanner PWA) Journey\n",
        "See **Journey 4 — Manual Agent PWA** in the P0 Exam-QR-Batch section for full button-by-button requirements.\n\n",
        "**Target app:** `apps/scanner-web` (port 4574) — **NOT BUILT**  \n",
        "**Backend seams:** `scanner` role, `requestUploadUrl`, `uploadAnswerSheets` — **EXISTS**\n\n",
        "### Summary Flow\n",
        "1. Login with school code + scanner credentials\n",
        "2. Select exam (published or grading status)\n",
        "3. Print QR batch per student roster\n",
        "4. Attach QR stickers to answer sheets\n",
        "5. Scan QR → resolve student + exam\n",
        "6. Capture multi-page answer sheet photos\n",
        "7. Submit via `requestUploadUrl` + `uploadAnswerSheets`\n",
        "8. End session with agent signature + QR audit\n",
        "\n\\newpage\n",
        "## Super-Admin Journey\n",
        admin.split("## B. Super-Admin")[1].strip() if "## B. Super-Admin" in admin else "",
        "\n\\newpage\n",
        "# P0 Exam-QR-Batch Journey\n",
        exam_body,
        "\n\\newpage\n",
        "# Data Model & Integrations\n",
        data_model if data_model else "_See EXAM-QR-BATCH-JOURNEY.md Data Model section._",
        "\n\n### AI Integration Points\n",
        cross.split("## AI surfaces")[1].split("## Access Denied")[0].strip()
        if "AI surfaces" in cross
        else "",
        "\n\\newpage\n",
        "# Gap Analysis\n",
        gap_analysis_section(),
        "\n\\newpage\n",
        "# Implementation Phases\n",
        implementation_phases(),
        "\n\\newpage\n",
        "# Appendix\n",
        appendix_section(),
        "\n\n### Marketing & Infrastructure Reference\n",
        infra,
    ]

    return "\n".join(parts)


def add_toc_field(doc: Document) -> None:
    """Insert Word TOC field (update fields in Word to populate)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    hint = doc.add_paragraph(
        "Right-click the table of contents and select Update Field → Update entire table."
    )
    hint.runs[0].italic = True
    hint.runs[0].font.size = Pt(9)
    hint.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_cover_page(final: Document) -> None:
    """Prepend professional cover page."""
    # Work on first section — insert paragraphs at start
    body = final.element.body
    # Build cover elements
    cover_doc = Document()
    for _ in range(6):
        cover_doc.add_paragraph("")
    title = cover_doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("LvlUp Product Requirements")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(0x1A, 0x56, 0x7A)

    subtitle = cover_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Complete Platform & P0 Exam-QR-Batch Journey")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    cover_doc.add_paragraph("")
    meta = cover_doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"Version {VERSION}",
        f"Date: {DOC_DATE}",
        "Document ID: LVLUP-PRODUCT-REQUIREMENTS",
        "Status: Production Requirements",
        "",
        "Auto-LevelUp · startup-mvp/lvlup",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(11)

    cover_doc.add_paragraph("")
    conf = cover_doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = conf.add_run("CONFIDENTIAL — Internal product requirements")
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    cover_doc.add_page_break()
    toc_heading = cover_doc.add_paragraph()
    th = toc_heading.add_run("Table of Contents")
    th.bold = True
    th.font.size = Pt(18)
    add_toc_field(cover_doc)
    cover_doc.add_page_break()

    # Prepend cover elements to final doc body
    cover_elements = list(cover_doc.element.body)
    for el in reversed(cover_elements):
        body.insert(0, el)


def merge_documents(cover_and_toc: Path, body: Path, output: Path) -> None:
    final = Document(str(body))
    add_cover_page(final)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    final.save(str(output))


def main() -> int:
    print("Building combined markdown...")
    combined = build_combined_markdown()
    TMP_MD.parent.mkdir(parents=True, exist_ok=True)
    TMP_MD.write_text(combined, encoding="utf-8")
    print(f"  Wrote {TMP_MD} ({len(combined):,} chars)")

    print("Converting to DOCX via pandoc...")
    extra_args = [
        "--from=markdown",
        "--to=docx",
        "--standalone",
        f"--metadata=title:LvlUp Product Requirements",
        f"--metadata=date:{DOC_DATE}",
    ]
    pypandoc.convert_file(
        str(TMP_MD),
        "docx",
        outputfile=str(TMP_DOCX),
        extra_args=extra_args,
    )
    print(f"  Body DOCX: {TMP_DOCX}")

    print("Adding cover page and TOC...")
    merge_documents(TMP_DOCX, TMP_DOCX, OUT_PATH)
    print(f"  Output: {OUT_PATH}")

    # Cleanup temp files
    for p in (TMP_MD, TMP_DOCX):
        try:
            p.unlink()
        except OSError:
            pass

    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
